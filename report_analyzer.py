import os
import base64
import tempfile
from PIL import Image
import fitz  # PyMuPDF
import docx
import io
import streamlit as st
from openai import OpenAI
from dotenv import load_dotenv  

load_dotenv()  
# Initialize OpenAI client
client = OpenAI()

# Function to encode image to base64
def encode_image(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
        temp_file.write(pdf_file.getvalue())
        temp_path = temp_file.name
    
    doc = fitz.open(temp_path)
    text = ""
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    
    # Clean up temp file
    os.unlink(temp_path)
    return text

# Function to extract text from DOCX
def extract_text_from_docx(docx_file):
    doc = docx.Document(io.BytesIO(docx_file.getvalue()))
    full_text = []
    
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

# Function to extract text from image using GPT-4o Vision
def extract_text_from_image_gpt4o(image_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_file:
        temp_file.write(image_file.getvalue())
        temp_path = temp_file.name
    
    # Encode the image to base64
    base64_image = encode_image(temp_path)
    
    # Clean up temp file
    os.unlink(temp_path)
    
    # Stage 1 prompt for text extraction
    extraction_prompt = """
    You are an expert document analyzer specializing in medical reports. Extract ALL text content from this document completely and accurately.
    
    Preserve the original structure and formatting as much as possible including:
    - Section headers and subheaders
    - Bullet points and numbered lists
    - Table structures (maintain rows and columns)
    - Paragraph breaks
    - Any emphasized text (bold, italics, underlined)
    
    If the document contains tables, recreate them using markdown table format.
    If the document contains any charts or graphs, describe what they show in [CHART DESCRIPTION] tags.
    If any text is illegible or uncertain, mark it with [?].
    If there are handwritten notes, extract them and mark with [HANDWRITTEN: text].
    Maintain all medical terminology exactly as written, including abbreviations, units, and values.
    
    IMPORTANT: Extract the COMPLETE document text without summarizing or omitting any content.
    """
    
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": extraction_prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                ]
            }
        ],
        max_tokens=4000
    )
    
    return response.choices[0].message.content

# Function to analyze medical report text
def analyze_medical_report(extracted_text):
    # Stage 2 prompt for medical analysis
    analysis_prompt = f"""
  

You are a medical data analyst. Analyze this medical report and provide a structured analysis using proper markdown formatting:

## 1. DISEASE IDENTIFICATION
- Diagnosed condition(s)
- Severity level observe from results
## 2. ABNORMAL MEASUREMENTS
- Values outside normal ranges (include reference ranges)
- Specific fluctuating values

## 3. RISK FACTORS
- Identified risk factors
- Lifestyle, genetic, and environmental factors

## 4. RECOMMENDED ACTIONS
- Provide key recommendations for improving the patient's health (e.g., lifestyle modifications, dietary changes, exercise, monitoring).

## 5. MEDICAL CONSULTATION
- Specialist referral needs


## 6. SUMMARY
- Key findings (2-3 points)
- Required next steps

IMPORTANT:
1. Use proper markdown formatting with headers (##) and bullet points (-)
2. Keep points concise (5-10 words when possible)
3. Use medical terminology from the report
4. Write "Not mentioned" for missing information
5. Format your response so it displays well in a web interface
6. Do not make assumptions beyond what's in the report but give suggestion for improvement after identified disease.


    Here is the extracted text to analyze:
    
    {extracted_text}
    """
    
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "user", "content": analysis_prompt}
        ],
        max_tokens=2000
    )
    
    return response.choices[0].message.content

# Streamlit app
def main():
    st.title("Medical Report Analyzer")
    st.write("Upload your medical report to extract and analyze important information.")
    
    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "jpg", "jpeg", "png"])
    
    if uploaded_file is not None:
        # Display file details
        file_details = {"Filename": uploaded_file.name, "File size": uploaded_file.size}
        st.write(file_details)
        
        # Extract text based on file type
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        with st.spinner("Extracting text from document..."):
            if file_type == "pdf":
                extracted_text = extract_text_from_pdf(uploaded_file)
            elif file_type == "docx":
                extracted_text = extract_text_from_docx(uploaded_file)
            elif file_type in ["jpg", "jpeg", "png"]:
                extracted_text = extract_text_from_image_gpt4o(uploaded_file)
            else:
                st.error("Unsupported file type")
                return
        
        # Display extracted text
        with st.expander("View Extracted Text", expanded=False):
            st.text_area("", extracted_text, height=400)
        
        # Analyze the extracted text
        if st.button("Analyze Medical Report"):
            with st.spinner("Analyzing the medical report..."):
                analysis_result = analyze_medical_report(extracted_text)
            
            st.subheader("Medical Report Analysis")
            st.markdown(analysis_result)
            
            # Option to download the analysis
            st.download_button(
                label="Download Analysis",
                data=analysis_result,
                file_name="medical_report_analysis.txt",
                mime="text/plain"
            )

if __name__ == "__main__":
    main()
